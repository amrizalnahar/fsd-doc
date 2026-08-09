#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Sisipkan logo brand ke halaman sampul .docx hasil konversi Pandoc.

Pandoc HANYA memakai gaya/margin/header-footer dari reference.docx — isi
body-nya (termasuk gambar) diabaikan. Jadi logo TIDAK BISA ditempel lewat
make-reference-docx.py; ia harus disisipkan sebagai post-process ke .docx
hasil konversi, tepat di atas paragraf bergaya "Title" yang dibuat Pandoc dari
metadata YAML `title:`.

Dipanggil otomatis oleh build-docx.ps1 sesudah Pandoc (bila brand.logo diisi
di doc-fsd.config.yml). brand.logo kosong/config tak ada -> dilewati tanpa
error (fitur opsional). brand.logo DIISI tetapi filenya tak ditemukan -> gagal
(exit 1) dengan pesan jelas, supaya salah ketik path tidak diam-diam membuat
sampul klien tanpa logo.

Jalankan manual:
    py insert-logo.py <file.docx>                       # cari config otomatis
    py insert-logo.py <file.docx> --logo <path/logo.png>
Butuh: pip install python-docx (+ pyyaml untuk baca config otomatis)
"""
import argparse
import os
import sys

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _find_config():
    d = os.getcwd()
    while True:
        for cand in (os.path.join(d, "docs", "tasks", "fsd", "doc-fsd.config.yml"),
                     os.path.join(d, "doc-fsd.config.yml")):
            if os.path.isfile(cand):
                return cand
        parent = os.path.dirname(d)
        if parent == d:
            return None
        d = parent


def _repo_root(config_path):
    """doc-fsd.config.yml biasanya di docs/tasks/fsd/ di dalam repo proyek;
    semua path config (termasuk brand.logo) relatif terhadap root repo itu,
    bukan terhadap lokasi file config sendiri."""
    norm = os.path.normpath(config_path).replace("\\", "/")
    suffix = "docs/tasks/fsd/doc-fsd.config.yml"
    if norm.endswith(suffix):
        return norm[: -len(suffix)].rstrip("/") or "/"
    return os.path.dirname(config_path)


def resolve_logo(explicit_logo):
    """Kembalikan (logo_path_or_None, config_path_or_None)."""
    if explicit_logo:
        return explicit_logo, None

    config_path = _find_config()
    if not config_path:
        return None, None

    try:
        import yaml
    except ImportError:
        print("[i] PyYAML belum terpasang - logo dilewati (pip install pyyaml).")
        return None, config_path

    try:
        with open(config_path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}
    except Exception as e:  # noqa: BLE001
        print(f"[i] Gagal membaca {config_path}: {e} - logo dilewati.")
        return None, config_path

    logo = (data.get("brand", {}) or {}).get("logo", "")
    if not logo:
        return None, config_path

    root = _repo_root(config_path)
    return os.path.join(root, logo), config_path


def insert_logo(docx_path, logo_path):
    doc = Document(docx_path)
    target = None
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Title":
            target = p
            break
    if target is None:
        print("[!] Paragraf sampul (gaya 'Title') tidak ditemukan - logo dilewati.")
        return False

    logo_p = target.insert_paragraph_before()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_p.add_run()
    run.add_picture(logo_path, height=Cm(2.5))
    doc.save(docx_path)
    return True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--logo", default=None, help="Path logo eksplisit (override config).")
    args = ap.parse_args()

    logo_path, _config_path = resolve_logo(args.logo)
    if not logo_path:
        # brand.logo kosong / config tak ada -> tanpa logo, ini BUKAN error.
        return 0

    if not os.path.isfile(logo_path):
        print(f"[x] brand.logo diatur ('{logo_path}') tetapi file tidak ditemukan.",
              file=sys.stderr)
        return 1

    if not os.path.isfile(args.docx):
        print(f"[x] File .docx tidak ditemukan: {args.docx}", file=sys.stderr)
        return 1

    if insert_logo(args.docx, logo_path):
        print(f"[ok] Logo disisipkan di sampul: {logo_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
