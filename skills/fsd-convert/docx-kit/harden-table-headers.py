#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Perbaiki portabilitas HEADER TABEL pada .docx hasil Pandoc.

Masalah: reference.docx memberi warna header (teks putih + latar brand) lewat
CONDITIONAL FORMATTING gaya tabel (w:tblStylePr type="firstRow"). Word menghormatinya,
tetapi Google Docs MENGABAIKAN tblStylePr saat impor — teks putih jatuh ke hitam
(sering di atas latar brand yang tetap gelap → kontras buruk / tak terbaca).

Solusi: bake teks putih + tebal + latar brand sebagai DIRECT FORMATTING langsung
pada sel & run baris header. Direct formatting dihormati Google Docs, LibreOffice,
dan Word — sehingga header tampil konsisten di mana pun dokumen dibuka.

Dipanggil otomatis oleh build-docx.ps1 SETELAH Pandoc. Mengubah file di tempat.

    py harden-table-headers.py <file.docx> [--reference <reference.docx>]

--reference : dipakai untuk mengambil warna latar header (fill firstRow) agar
              sama persis dengan brand di reference.docx. Bila tidak diberikan
              atau gagal dibaca, dipakai default netral.

Butuh: pip install python-docx   (bila belum ada → skrip skip dengan pesan; .docx
       tetap valid & benar di Word, hanya belum di-harden untuk Google Docs).
"""

import sys

WHITE_HEX = "FFFFFF"
DEFAULT_PRIMARY = "1F5FA8"   # selaras default netral make-reference-docx.py


def _parse_args(argv):
    docx_path, ref_path = None, None
    i = 0
    while i < len(argv):
        a = argv[i]
        if a in ("--reference", "-r"):
            i += 1
            ref_path = argv[i] if i < len(argv) else None
        elif docx_path is None:
            docx_path = a
        i += 1
    return docx_path, ref_path


def read_primary_fill(reference_path, default=DEFAULT_PRIMARY):
    """Ambil warna latar header (fill firstRow) dari gaya 'Table' reference.docx."""
    if not reference_path:
        return default
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return default
    try:
        ref = Document(reference_path)
    except Exception:                                   # noqa: BLE001
        return default
    for st in ref.styles.element.findall(qn("w:style")):
        if st.get(qn("w:styleId")) != "Table":
            continue
        for sp in st.findall(qn("w:tblStylePr")):
            if sp.get(qn("w:type")) != "firstRow":
                continue
            tcPr = sp.find(qn("w:tcPr"))
            shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
            fill = shd.get(qn("w:fill")) if shd is not None else None
            if fill and fill.lower() != "auto":
                return fill.upper()
    return default


def _is_header_row(row, qn):
    trPr = row._tr.find(qn("w:trPr"))
    return trPr is not None and trPr.find(qn("w:tblHeader")) is not None


def _set_cell_shading(cell, fill, qn, OxmlElement):
    """Set latar sel sebagai direct formatting, sisipkan w:shd pada posisi sah."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    # w:shd harus berada sesudah elemen-elemen ini (skema CT_TcPr) dan sebelum sisanya.
    pre = {qn(x) for x in (
        "w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders")}
    anchor = None
    for child in tcPr:
        if child.tag not in pre:
            anchor = child
            break
    if anchor is not None:
        anchor.addprevious(shd)
    else:
        tcPr.append(shd)


def harden(docx_path, reference_path):
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("[skip] python-docx belum terpasang (pip install python-docx) — "
              "header .docx tidak di-harden untuk Google Docs (tetap benar di Word).")
        return 0

    primary = read_primary_fill(reference_path)
    white = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document(docx_path)
    n_tables = 0
    n_cells = 0
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        header_rows = [r for r in rows if _is_header_row(r, qn)]
        if not header_rows:                 # tabel tanpa penanda tblHeader → pakai baris pertama
            header_rows = [rows[0]]
        touched = False
        for row in header_rows:
            for cell in row.cells:
                _set_cell_shading(cell, primary, qn, OxmlElement)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = white
                        run.font.bold = True
                n_cells += 1
                touched = True
        if touched:
            n_tables += 1

    doc.save(docx_path)
    print(f"[ok] Header tabel di-harden untuk Google Docs: "
          f"{n_tables} tabel, {n_cells} sel (teks #{WHITE_HEX} + latar #{primary}).")
    return 0


def main(argv):
    docx_path, ref_path = _parse_args(argv)
    if not docx_path:
        print("Pemakaian: py harden-table-headers.py <file.docx> [--reference <reference.docx>]")
        return 2
    import os
    if not os.path.isfile(docx_path):
        print(f"[x] File tidak ditemukan: {docx_path}")
        return 1
    return harden(docx_path, ref_path)


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
