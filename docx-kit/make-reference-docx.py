#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Pembuat reference.docx (template gaya Word) untuk FSD.

Dipakai oleh Pandoc saat mengonversi fsd-*.md menjadi .docx:
    pandoc fsd.md -o fsd.docx --reference-doc=reference.docx --toc ...

Pandoc HANYA mengambil DEFINISI GAYA (styles), margin halaman, serta
header/footer dari file ini — isi body-nya diabaikan. Jadi cukup ubah gaya
di sini lalu regenerate: hasilnya otomatis dipakai semua dokumen FSD.

Brand (warna, font, label header) diambil dari doc-fsd.config.yml proyek:
    brand.color_primary       -> heading & tabel header
    brand.color_primary_dark  -> judul sampul
    brand.color_accent        -> aksen sekunder (opsional)
    brand.font_heading / font_body
    project.name + project.client -> label header dokumen
Bila config atau PyYAML tak ada, dipakai default netral di bawah.

Jalankan:  py make-reference-docx.py                 # cari config otomatis
           DOC_FSD_CONFIG=/path/config.yml py make-reference-docx.py
Butuh   :  pip install python-docx   (pip install pyyaml untuk baca config)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------- config loader
def _hex_to_rgb(h):
    h = str(h).lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _find_config():
    """Cari doc-fsd.config.yml: env DOC_FSD_CONFIG, lalu naik dari cwd."""
    env = os.environ.get("DOC_FSD_CONFIG")
    if env and os.path.isfile(env):
        return env
    d = os.getcwd()
    while True:
        for cand in (os.path.join(d, "doc-fsd.config.yml"),
                     os.path.join(d, "docs", "tasks", "fsd", "doc-fsd.config.yml")):
            if os.path.isfile(cand):
                return cand
        parent = os.path.dirname(d)
        if parent == d:
            return None
        d = parent


def load_brand():
    """Kembalikan dict brand efektif (config bila ada, selain itu default netral)."""
    # default netral (bukan brand klien tertentu)
    cfg = {
        "color_primary": "1F5FA8",
        "color_primary_dark": "133C6B",
        "color_accent": "2E7D46",
        "font_heading": "Georgia",
        "font_body": "Calibri",
        "doc_label": "Functional Specification Document",
    }
    path = _find_config()
    if not path:
        print("[i] doc-fsd.config.yml tidak ditemukan — pakai brand default netral.")
        return cfg
    try:
        import yaml
        with open(path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f) or {}
    except ImportError:
        print("[i] PyYAML belum terpasang (pip install pyyaml) — pakai default.")
        return cfg
    except Exception as e:                      # noqa: BLE001
        print(f"[i] Gagal membaca {path}: {e} — pakai default.")
        return cfg
    brand = data.get("brand", {}) or {}
    project = data.get("project", {}) or {}
    for k in ("color_primary", "color_primary_dark", "color_accent",
              "font_heading", "font_body"):
        if brand.get(k):
            cfg[k] = brand[k]
    name = project.get("name")
    client = project.get("client")
    label = "Functional Specification Document"
    if name:
        label = f"{label} — {name}" + (f" ({client})" if client else "")
    cfg["doc_label"] = label
    print(f"[i] Brand dari config: {path}")
    return cfg


_BRAND = load_brand()

# ---------------------------------------------------------------- brand palette
PRIMARY      = _hex_to_rgb(_BRAND["color_primary"])       # heading, judul
PRIMARY_DARK = _hex_to_rgb(_BRAND["color_primary_dark"])  # judul sampul
ACCENT       = _hex_to_rgb(_BRAND["color_accent"])        # aksen sekunder opsional
DARK         = RGBColor(0x1F, 0x2D, 0x3D)   # #1F2D3D — teks isi
GREY         = RGBColor(0x5A, 0x6A, 0x7A)   # abu-abu — caption, footer, subjudul

PRIMARY_HEX = str(_BRAND["color_primary"]).lstrip("#").upper()
BORDER_HEX  = "BFBFBF"

HEADING_FONT = _BRAND["font_heading"]   # serif elegan untuk judul & heading
BODY_FONT    = _BRAND["font_body"]      # sans bersih untuk isi
DOC_LABEL    = _BRAND["doc_label"]      # label header tiap halaman


# ---------------------------------------------------------------- helpers
def ensure_para_style(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_font(style, name=None, size=None, bold=None, color=None, italic=None):
    f = style.font
    if name is not None:
        f.name = name
        # pastikan berlaku juga untuk east-asian / complex script
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), name)
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color


def set_spacing(style, before=None, after=None, line=None,
                keep_with_next=None, page_break_before=None, align=None):
    pf = style.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next
    if page_break_before is not None:
        pf.page_break_before = page_break_before
    if align is not None:
        pf.alignment = align


def add_field(paragraph, field, size=9, color=GREY):
    """Sisipkan field Word (mis. PAGE / NUMPAGES) ke sebuah paragraf."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color


def style_table(doc):
    """Gaya tabel 'Table' (dipakai Pandoc): grid tipis + header berwarna brand."""
    try:
        st = doc.styles["Table"]
    except KeyError:
        st = doc.styles.add_style("Table", WD_STYLE_TYPE.TABLE)
    set_font(st, name=BODY_FONT, size=10, color=DARK)

    el = st.element

    # paragraf di dalam sel: rata KIRI (mengalahkan Normal yang justify),
    # dipasang di level style tabel sebagai jaring pengaman selain gaya "Compact".
    p_pr = OxmlElement("w:pPr")
    p_jc = OxmlElement("w:jc"); p_jc.set(qn("w:val"), "left")
    p_pr.append(p_jc)
    rpr_el = el.find(qn("w:rPr"))
    if rpr_el is not None:
        rpr_el.addprevious(p_pr)
    else:
        el.append(p_pr)

    tblPr = el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        el.append(tblPr)

    # blok tabel dipusatkan di halaman (rata tengah). jc harus mendahului
    # tblCellMar/tblBorders sesuai skema OOXML → sisipkan sebagai anak pertama.
    tbl_jc = OxmlElement("w:jc"); tbl_jc.set(qn("w:val"), "center")
    tblPr.insert(0, tbl_jc)

    # padding sel supaya tidak sesak
    margins = OxmlElement("w:tblCellMar")
    for edge, w in (("top", 40), ("bottom", 40), ("start", 80), ("end", 80)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(w)); m.set(qn("w:type"), "dxa")
        margins.append(m)
    tblPr.append(margins)

    # garis grid tipis abu-abu
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")            # 0.5pt
        b.set(qn("w:color"), BORDER_HEX)
        borders.append(b)
    tblPr.append(borders)

    # baris header: latar warna brand + teks putih tebal
    first = OxmlElement("w:tblStylePr")
    first.set(qn("w:type"), "firstRow")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b"); rPr.append(b)
    c = OxmlElement("w:color"); c.set(qn("w:val"), "FFFFFF"); rPr.append(c)
    first.append(rPr)
    tcPr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), PRIMARY_HEX)
    tcPr.append(shd)
    first.append(tcPr)
    el.append(first)


def add_header_footer(doc):
    for section in doc.sections:
        # ---- footer: nomor halaman terpusat
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Halaman ")
        r.font.name = BODY_FONT; r.font.size = Pt(9); r.font.color.rgb = GREY
        add_field(fp, "PAGE")
        r2 = fp.add_run(" dari ")
        r2.font.name = BODY_FONT; r2.font.size = Pt(9); r2.font.color.rgb = GREY
        add_field(fp, "NUMPAGES")

        # ---- header: label ringkas rata kanan
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(DOC_LABEL)
        hr.font.name = BODY_FONT; hr.font.size = Pt(8); hr.font.color.rgb = GREY
        hr.italic = True


# ---------------------------------------------------------------- build
def build(output_path):
    doc = Document()

    # margin halaman
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- Normal / isi
    normal = doc.styles["Normal"]
    set_font(normal, name=BODY_FONT, size=11, color=DARK)
    set_spacing(normal, after=6, line=1.15, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- Heading 1..4  (BAB & sub-bab). Penomoran sudah manual di .md,
    #      jadi JANGAN pakai --number-sections di Pandoc.
    h1 = doc.styles["Heading 1"]
    set_font(h1, name=HEADING_FONT, size=18, bold=True, color=PRIMARY)
    set_spacing(h1, before=0, after=10, keep_with_next=True, page_break_before=True)

    h2 = doc.styles["Heading 2"]
    set_font(h2, name=HEADING_FONT, size=14, bold=True, color=PRIMARY)
    set_spacing(h2, before=14, after=6, keep_with_next=True)

    h3 = doc.styles["Heading 3"]
    set_font(h3, name=HEADING_FONT, size=12, bold=True, color=DARK)
    set_spacing(h3, before=10, after=4, keep_with_next=True)

    h4 = doc.styles["Heading 4"]
    set_font(h4, name=BODY_FONT, size=11, bold=True, color=PRIMARY)
    set_spacing(h4, before=8, after=2, keep_with_next=True)

    # ---- Sampul: Title / Subtitle / Author / Date
    title = doc.styles["Title"]
    set_font(title, name=HEADING_FONT, size=30, bold=True, color=PRIMARY_DARK)
    set_spacing(title, before=150, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    subtitle = ensure_para_style(doc, "Subtitle")
    set_font(subtitle, name=BODY_FONT, size=16, bold=False, color=GREY)
    set_spacing(subtitle, before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)

    author = ensure_para_style(doc, "Author")
    set_font(author, name=BODY_FONT, size=12, color=DARK)
    set_spacing(author, before=24, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

    date = ensure_para_style(doc, "Date")
    set_font(date, name=BODY_FONT, size=12, color=GREY)
    set_spacing(date, before=0, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- TOC: mulai di halaman baru → sampul berdiri sendiri di halaman 1
    toc = ensure_para_style(doc, "TOC Heading")
    set_font(toc, name=HEADING_FONT, size=16, bold=True, color=PRIMARY)
    set_spacing(toc, before=0, after=8, page_break_before=True)

    # ---- Gambar (figure): Pandoc memberi paragraf gambar gaya "Figure".
    #      Dipusatkan agar gambar rata tengah di halaman.
    fig = ensure_para_style(doc, "Figure")
    set_font(fig, name=BODY_FONT, size=11, color=DARK)
    set_spacing(fig, before=6, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- Caption gambar
    cap = ensure_para_style(doc, "Image Caption")
    set_font(cap, name=BODY_FONT, size=9, italic=True, color=GREY)
    set_spacing(cap, before=2, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- Compact (list rapat + isi sel tabel yang dipakai Pandoc).
    #      Rata KIRI: mengalahkan Normal (justify) untuk teks di dalam sel tabel.
    compact = ensure_para_style(doc, "Compact")
    set_font(compact, name=BODY_FONT, size=11, color=DARK)
    set_spacing(compact, after=2, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ---- Quote / Block Text
    try:
        quote = doc.styles["Quote"]
    except KeyError:
        quote = ensure_para_style(doc, "Quote")
    set_font(quote, name=BODY_FONT, size=11, italic=True, color=GREY)
    set_spacing(quote, before=6, after=6)

    # ---- tabel & header/footer
    style_table(doc)
    add_header_footer(doc)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reference.docx")
    path = build(out)
    print(f"reference.docx dibuat: {path}")
